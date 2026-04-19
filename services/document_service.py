import os
import re
import logging
import aiofiles
from config import UPLOADS_DIR, MAX_FILE_SIZE_MB

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def _safe_dirname(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", name)
    name = re.sub(r"\s+", "_", name)
    name = name.strip("._")
    return name[:60] or "course"


def get_upload_path(user_id: int, course_id: str, filename: str, course_name: str = "") -> str:
    if course_name:
        folder = f"{_safe_dirname(course_name)}_{course_id[:8]}"
    else:
        folder = course_id
    directory = os.path.join(UPLOADS_DIR, str(user_id), folder)
    try:
        os.makedirs(directory, exist_ok=True)
    except OSError as e:
        logger.error(f"Cannot create upload dir '{directory}': {e}")
        fallback = os.path.join(UPLOADS_DIR, str(user_id), course_id)
        os.makedirs(fallback, exist_ok=True)
        directory = fallback
    return os.path.join(directory, filename)


async def save_file(file_bytes: bytes, user_id: int, course_id: str, filename: str, course_name: str = "") -> str:
    path = get_upload_path(user_id, course_id, filename, course_name)
    try:
        async with aiofiles.open(path, "wb") as f:
            await f.write(file_bytes)
        logger.info(f"Saved upload: {path} ({len(file_bytes):,} bytes)")
        return path
    except OSError as e:
        logger.error(f"Failed to write file '{path}': {e}")
        raise RuntimeError(f"خطا در ذخیره فایل روی دیسک: {e}") from e


def extract_text_from_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[صفحه {i + 1}]\n{text}")
        return "\n\n".join(pages)
    except Exception as e:
        logger.error(f"PDF extraction error for '{path}': {e}")
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error for '{path}': {e}")
        return ""


def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.error(f"TXT extraction error for '{path}': {e}")
        return ""


def extract_text(path: str, file_type: str) -> str:
    ext = file_type.lower()
    if ext == "pdf":
        return extract_text_from_pdf(path)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(path)
    elif ext == "txt":
        return extract_text_from_txt(path)
    return ""


def get_file_type(filename: str) -> str | None:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"
    if ext == "txt":
        return "txt"
    return None


def is_valid_file_size(size_bytes: int) -> bool:
    return size_bytes <= MAX_FILE_SIZE_BYTES


def delete_file(path: str):
    try:
        if path and os.path.exists(path):
            os.remove(path)
            logger.info(f"Deleted: {path}")
    except Exception as e:
        logger.error(f"File delete error for '{path}': {e}")
